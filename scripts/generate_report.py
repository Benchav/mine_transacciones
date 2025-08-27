# scripts/generate_report.py
from docx import Document
from docx.shared import Inches
import pandas as pd, os

doc = Document()
doc.add_heading('Informe - Minería de Transacciones', level=1)

# Preprocesado
orig = pd.read_csv('data/ventas_ejemplo.csv', encoding='ISO-8859-1').shape[0]
clean = pd.read_csv('outputs/ventas_preprocessed.csv').shape[0]
doc.add_heading('Preprocesamiento', level=2)
doc.add_paragraph(f'Registros originales: {orig}\nRegistros después limpieza: {clean}')

# Reglas top
doc.add_heading('Top reglas (por lift)', level=2)
if os.path.exists('outputs_apriori/rules_apriori.csv'):
    rules = pd.read_csv('outputs_apriori/rules_apriori.csv').head(10)
    doc.add_paragraph(rules[['antecedents','consequents','support','confidence','lift']].to_string(index=False))
else:
    doc.add_paragraph('No se generaron reglas Apriori. Revisa outputs_apriori/rules_apriori.csv')

# Gráficos
doc.add_heading('Visualizaciones', level=2)
if os.path.exists('outputs/top10_support.png'):
    doc.add_paragraph('Top items por support')
    doc.add_picture('outputs/top10_support.png', width=Inches(5))
if os.path.exists('outputs_apriori/rules_network.png'):
    doc.add_paragraph('Red de reglas')
    doc.add_picture('outputs_apriori/rules_network.png', width=Inches(5))
if os.path.exists('outputs/rfm_radar.png'):
    doc.add_paragraph('Radar de perfiles RFM')
    doc.add_picture('outputs/rfm_radar.png', width=Inches(5))

# Clusters
doc.add_heading('RFM y Clusters', level=2)
if os.path.exists('outputs/Perfil_de_clusters.csv'):
    profile = pd.read_csv('outputs/Perfil_de_clusters.csv')
    doc.add_paragraph(profile.to_string(index=False))

doc.save('outputs/informe_minero.docx')
print('Informe guardado en outputs/informe_minero.docx')