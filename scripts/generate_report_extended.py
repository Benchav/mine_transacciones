# scripts/generate_report_extended.py
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

OUTDIR = 'outputs'
APRIORI_DIR = 'outputs_apriori'
DOC_PATH = os.path.join(OUTDIR, 'informe_minero_extended.docx')

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_table_from_df(doc, df, title=None, col_widths=None):
    if title:
        doc.add_paragraph(title, style='Intense Quote')
    # create table (including header)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Light List Accent 1'
    hdr_cells = table.rows[0].cells
    for i, c in enumerate(df.columns):
        hdr_cells[i].text = str(c)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(df.columns):
            val = row[c]
            if pd.isna(val):
                cells[i].text = ''
            else:
                cells[i].text = str(val)
    doc.add_paragraph('')  # spacer

def try_add_image(doc, path, caption=None, width_inches=5.5):
    if os.path.exists(path):
        try:
            doc.add_paragraph(caption if caption else '')
            doc.add_picture(path, width=Inches(width_inches))
            last_par = doc.paragraphs[-1]
            last_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph('')
            return True
        except Exception as e:
            doc.add_paragraph(f"(No se pudo insertar la imagen {path}: {e})")
    else:
        doc.add_paragraph(f"(Imagen no encontrada: {path})")
    return False

def main():
    os.makedirs(OUTDIR, exist_ok=True)
    doc = Document()
    doc.add_heading('Informe extendido - Minería de Transacciones', level=0)
    doc.add_paragraph('Proyecto: Market Basket Analysis y Segmentación RFM')
    doc.add_paragraph('Creado por: Joshua Chaves — https://joshuachavez.vercel.app/')
    doc.add_paragraph('')

    # 1) Preprocesado - contar filas
    preproc_path = os.path.join(OUTDIR, 'ventas_preprocessed.csv')
    if os.path.exists(preproc_path):
        dfp = pd.read_csv(preproc_path, parse_dates=['InvoiceDate'], dayfirst=True, encoding='ISO-8859-1')
        doc.add_heading('Preprocesamiento', level=1)
        doc.add_paragraph(f'Registros procesados: {len(dfp)}')
        doc.add_paragraph(f'Fechas (min, max): {dfp["InvoiceDate"].min()} — {dfp["InvoiceDate"].max()}')
    else:
        doc.add_heading('Preprocesamiento', level=1)
        doc.add_paragraph('No se encontró outputs/ventas_preprocessed.csv')

    # 2) Reglas: apriori o fallback
    doc.add_heading('Reglas de asociación (Top reglas)', level=1)
    rules_path_ap = os.path.join(APRIORI_DIR, 'rules_apriori.csv')
    rules_path_fb = os.path.join(OUTDIR, 'rules_pairs_top.csv')
    rules_df = None
    used = ''
    if os.path.exists(rules_path_ap) and os.path.getsize(rules_path_ap) > 0:
        try:
            rules_df = pd.read_csv(rules_path_ap)
            used = 'Apriori'
        except:
            rules_df = None
    if rules_df is None and os.path.exists(rules_path_fb):
        try:
            rules_df = pd.read_csv(rules_path_fb)
            used = 'Fallback (pares)'
        except:
            rules_df = None

    if rules_df is None or rules_df.empty:
        doc.add_paragraph('No se encontraron reglas en outputs_apriori ni en outputs/rules_pairs_top.csv')
    else:
        doc.add_paragraph(f'Se usan reglas desde: {used}')
        # seleccionar columnas a mostrar
        # adaptar a nombres posibles
        display_cols = []
        if 'antecedents' in rules_df.columns and 'consequents' in rules_df.columns:
            display_cols = ['antecedents','consequents','support','confidence','lift'] if 'confidence' in rules_df.columns else ['antecedents','consequents','support','lift']
        elif 'antecedent' in rules_df.columns and 'consequent' in rules_df.columns:
            display_cols = ['antecedent','consequent','support','confidence_a_b','lift']
        # preparar top10 por lift si existe columna
        if 'lift' in rules_df.columns:
            top_rules = rules_df.sort_values(by='lift', ascending=False).head(10)
        else:
            top_rules = rules_df.head(10)
        # simplificar sets si es necesario (si hay listas o frozenset)
        def clean_val(v):
            if pd.isna(v):
                return ''
            if isinstance(v, str):
                return v
            try:
                return str(v)
            except:
                return ''
        top_rules = top_rules[display_cols].copy()
        top_rules = top_rules.applymap(clean_val)
        add_table_from_df(doc, top_rules, title='Top 10 reglas (ordenadas por lift)')
        doc.add_paragraph('')

    # 3) Interpretaciones (si existe file)
    interp_path = os.path.join(OUTDIR, 'rules_interpretations.txt')
    if os.path.exists(interp_path):
        doc.add_heading('Interpretaciones (Top 5)', level=1)
        with open(interp_path, 'r', encoding='utf-8') as f:
            text = f.read().strip()
        for block in text.split('\n\n'):
            if block.strip():
                doc.add_paragraph(block)
    else:
        doc.add_paragraph('No se encontró outputs/rules_interpretations.txt')

    # 4) Perfil de clusters
    profile_path = os.path.join(OUTDIR, 'Perfil_de_clusters.csv')
    if os.path.exists(profile_path):
        doc.add_heading('Perfil de Clusters (RFM)', level=1)
        prof = pd.read_csv(profile_path)
        add_table_from_df(doc, prof, title='Perfil de clusters (medias / medianas / count)')
    else:
        doc.add_paragraph('No se encontró outputs/Perfil_de_clusters.csv')

    # 5) Insertar figuras disponibles
    doc.add_heading('Gráficos', level=1)
    try_add_image(doc, os.path.join(OUTDIR, 'top10_support.png'), caption='Top 10 items por support')
    # preferir apriori network si existe, sino fallback
    if os.path.exists(os.path.join(APRIORI_DIR, 'rules_network.png')):
        try_add_image(doc, os.path.join(APRIORI_DIR, 'rules_network.png'), caption='Red de reglas (Apriori)')
    elif os.path.exists(os.path.join(OUTDIR, 'rules_network_fallback.png')):
        try_add_image(doc, os.path.join(OUTDIR, 'rules_network_fallback.png'), caption='Red de reglas (fallback)')
    elif os.path.exists(os.path.join(OUTDIR, 'rules_network.png')):
        try_add_image(doc, os.path.join(OUTDIR, 'rules_network.png'), caption='Red de reglas')

    try_add_image(doc, os.path.join(OUTDIR, 'rfm_radar.png'), caption='Radar perfil por cluster')
    try_add_image(doc, os.path.join(OUTDIR, 'rfm_scatter_colored.png'), caption='Scatter Recency vs Monetary coloreado por cluster')
    try_add_image(doc, os.path.join(OUTDIR, 'rfm_scatter.png'), caption='Scatter Recency vs Monetary (sin color)')

    # 6) Conclusiones y recomendaciones (plantilla)
    doc.add_heading('Conclusiones y recomendaciones', level=1)
    doc.add_paragraph('Resumen de acciones sugeridas (ejemplos):')
    doc.add_paragraph('- Priorizar reglas con lift alto y soporte razonable (>0.5%) para campañas de cross-sell.')
    doc.add_paragraph('- Ejecutar pruebas A/B para las 3 reglas con mayor lift y soporte mayor (ver Top 10).')
    doc.add_paragraph('- Para segmentation: campañas de reenganche para clusters con recency alto; upsell para clusters con monetary alto.')

    # 7) Guardar
    doc.save(DOC_PATH)
    print('Informe extendido guardado en:', DOC_PATH)

if __name__ == "__main__":
    main()
